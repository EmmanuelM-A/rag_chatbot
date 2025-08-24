"""
Responsible for loading documents of different formats.
"""
import os
from abc import abstractmethod, ABC
from pathlib import Path

import docx
import fitz
import zipfile

from docx.opc.exceptions import PackageNotFoundError, OpcError

from src.components.ingestion.document import FileDocument, \
    FileDocumentMetadata
from src.utils import exceptions


class DocumentLoader(ABC):
    """
    Abstract base class for document loaders.
    This class defines the interface for loading data of various file
    formats.
    """

    @abstractmethod
    def read_data(self, file_path: str) -> str:
        """
        Reads data from the specified file path.

        :param file_path: The file path to read from.
        :return: The contents of the file.
        """

        raise NotImplementedError("Subclasses must implement this method.")

    def load_data(self, file_path: str) -> FileDocument:
        """
        Load data from the specified file path and convert it to a FileDocument.

        :param file_path: The file path from which to load data (e.g., file path, URL).
        :return: The FileDocument representation of the file.
        """

        path = Path(file_path)

        if not path.exists():
            raise exceptions.FileDoesNotExist(
                path=file_path
            )

        if not path.is_file():
            raise exceptions.FileDoesNotExist(
                path=file_path,
                reason="The file path is not a file."
            )

        try:
            content = self.read_data(file_path)
        except (OSError, ValueError, Exception) as e:
            raise exceptions.DocumentLoadError(
                document_path=file_path,
                original_error=e
            )

        filename = os.path.splitext(file_path)[0].lower()

        file_extension = os.path.splitext(file_path)[1].lower()

        file_metadata = FileDocumentMetadata(
            filename=filename,
            file_extension=file_extension,
            author=None,
            source=file_path
        )

        return FileDocument(
            content=content,
            metadata=file_metadata
        )

    def __str__(self):
        """
        String representation of the DocumentLoader.
        :return: A string indicating the type of document loader.
        """

        return f"{self.__class__.__name__}"


class TxTDocumentLoader(DocumentLoader):
    """
    Document loader responsible for loading txt files.
    """

    def read_data(self, file_path: str) -> str:
        try:
            with open(file_path, "r", encoding="utf-8") as file:
                content = file.read()
        except UnicodeDecodeError:
            with open(file_path, 'r', encoding='latin-1') as file:
                content = file.read()
        except (FileNotFoundError, PermissionError, IsADirectoryError, OSError) as e:
            raise e

        return content


class MarkdownDocumentLoader(DocumentLoader):
    """
    Document loader responsible for loading markdown files.
    """

    def read_data(self, file_path: str) -> str:
        try:
            with open(file_path, "r", encoding="utf-8") as file:
                content = file.read()
        except UnicodeDecodeError:
            with open(file_path, 'r', encoding='latin-1') as file:
                content = file.read()
        except (FileNotFoundError, PermissionError, IsADirectoryError, OSError) as e:
            raise e

        return content


class PDFDocumentLoader(DocumentLoader):
    """
    Document loader responsible for loading pdf files.
    """

    def read_data(self, file_path: str) -> str:
        try:
            content = ""
            with fitz.open(file_path) as doc:
                for page in doc:
                    content += page.get_text("text")

            return content
        except fitz.FileNotFoundError as e:
            # PDF file doesn't exist
            raise e
        except fitz.EmptyFileError as e:
            # PDF file is empty
            raise e
        except fitz.FileDataError as e:
            # PDF file is corrupted or invalid format
            raise e
        except PermissionError as e:
            # No permission to read the file
            raise e
        except MemoryError as e:
            # PDF file is too large to process
            raise e
        except Exception as e:
            # Any other unexpected errors
            raise e


class DocxDocumentLoader(DocumentLoader):
    """
    Document loader responsible for loading docx files.
    """

    def read_data(self, file_path: str) -> str:
        try:
            doc = docx.Document(str(file_path))
            content = "\n".join(
                [para.text for para in doc.paragraphs]
            )
            return content
        except FileNotFoundError as e:
            # DOCX file doesn't exist
            raise e
        except PackageNotFoundError as e:
            # Invalid DOCX file or wrong format
            raise e
        except OpcError as e:
            # DOCX file is corrupted
            raise e
        except PermissionError as e:
            # No permission to read the file
            raise e
        except zipfile.BadZipFile as e:
            # DOCX is essentially a ZIP file, this catches ZIP-related errors
            raise e
        except MemoryError as e:
            # DOCX file is too large to process
            raise e
        except Exception as e:
            # Any other unexpected errors
            raise e
