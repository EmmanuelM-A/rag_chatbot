import os
from utils.constants import ALLOWED_FILE_EXTENSIONS

import docx
import PyPDF2


class FileDocument:
    def __init__(self, content, metadata):
        """
        Represents a document with its content and metadata.

        :param content: The main content of the document (e.g., text from a file).
        :param metadata: A dictionary or object containing metadata about the document (e.g., filename, file extension,
        author, etc.).
        """
        self.content = content
        self.metadata = metadata


def load_documents(path_to_directory):
    """
    Reads all files in the directory that match the allowed extensions and converts them into Document objects.

    :param path_to_directory: The path to the folder containing the document files.
    :return: A list of loaded Document objects.
    """

    documents = []

    # Iterate through all files in the directory and subdirectories
    for root, _, files in os.walk(path_to_directory):
        for file in files:
            # Get the file extension (in lowercase)
            ext = os.path.splitext(file)[1].lower()

            if ext not in ALLOWED_FILE_EXTENSIONS:
                print(f"The file: {file} is not allowed!")
                continue

            file_path = os.path.join(root, file)

            try:
                # Handle plain text and markdown files
                if ext in [".txt", ".md"]:
                    with open(file_path, 'r', encoding="utf-8") as f:
                        content = f.read()

                # Handle PDF files using PyPDF2
                elif ext == ".pdf":
                    content = ""
                    with open(file_path, 'rb') as f:
                        reader = PyPDF2.PdfReader(f)
                        # Extract text from all pages
                        for page in reader.pages:
                            content += page.extract_text() or ""

                # Handles Word (.docx) files using python-docx
                elif ext == ".docx":
                    doc = docx.Document(str(file_path))

                    # Combine all paragraph texts into a single string
                    content = "\n".join([para.text for para in doc.paragraphs])

                else:
                    continue # Skip unknown file types

                # Add the document to the list with metadata
                documents.append(FileDocument(content, {'source': file_path}))

            except Exception as e:
                # Log an error if reading fails
                print(f"Failed to read {file_path}: {e}")

    return documents
