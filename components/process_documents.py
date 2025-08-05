"""
Handles document ingestion, cleaning, splitting, and preparation of source
documents before embedding and storage.
"""

import os
from dataclasses import dataclass
from typing import Any

import docx
import fitz
from langchain_text_splitters import RecursiveCharacterTextSplitter
from sqlalchemy.testing.suite.test_reflection import metadata

from components.config import ALLOWED_FILE_EXTENSIONS, CHUNK_SIZE, CHUNK_OVERLAP
from utils.logger import get_logger

logger = get_logger(__name__)

from dataclasses import dataclass
from typing import Optional
from datetime import datetime


@dataclass
class FileDocumentMetadata:
    """
    Represents metadata associated with a file-based document.

    This class is used to capture contextual information about the
    document, which can help with filtering, tracing source documents,
    or debugging during document processing in a RAG pipeline.

    Attributes:
        filename (str): The original name of the file.
        file_extension (str): The file's extension (e.g., '.pdf', '.txt').
        author (Optional[str]): The author of the document, if known.
        created_at (Optional[datetime]): Timestamp when the document was created.
        modified_at (Optional[datetime]): Timestamp when the document was last modified.
        source (Optional[str]): Where the file came from (e.g., URL, uploader).
        document_id (Optional[str]): An internal unique ID, if available.
    """

    filename: str
    file_extension: str
    author: Optional[str] = None
    created_at: Optional[datetime] = None
    modified_at: Optional[datetime] = None
    source: Optional[str] = None
    document_id: Optional[str] = None


@dataclass
class FileDocument:
    """
    Represents a document with its content and metadata.

    It contains the main content of the document (e.g., text from a file) and
    a dictionary or object containing metadata about the document (e.g.,
    filename, file extension, author, etc.).
    """

    content: str
    metadata: FileDocumentMetadata

    def __str__(self):
        """Returns the name of the document"""

        return self.metadata.filename



class DocumentProcessor:
    """
    Prepares raw documents for indexing into the vector store.
    """

    def __init__(self, path_to_directory: str) -> None:
        self.path_to_directory = path_to_directory

    def load_documents(self):
        pass

    def _clean_documents(self):
        pass

    def chunk_documents(self):
        pass

    def process_documents(self):
        pass

def load_documents(path_to_directory):
    """
    Reads all files in the directory that match the allowed extensions and converts them into Document objects.

    :param path_to_directory: The path to the folder containing the document files.

    :returns list[FileDocument]: A list of loaded FileDocument objects
    """

    if not os.path.exists(path_to_directory):
        logger.error(f"The directory {path_to_directory} does not exist!")
        return []

    if not os.path.isdir(path_to_directory):
        logger.error(f"The path {path_to_directory} is not a directory!")
        return []

    logger.debug(f"Loading documents from: {path_to_directory}.")

    documents = []

    # Iterate through all files in the directory and subdirectories
    for root, _, files in os.walk(path_to_directory):
        for file in files:
            # Get the file extension (in lowercase)
            ext = os.path.splitext(file)[1].lower()

            if ext not in ALLOWED_FILE_EXTENSIONS:
                logger.warning(f"Skipping unsupported file: {file}.")
                continue

            file_path = os.path.join(root, file)

            logger.debug(f"Reading the file: {file_path}")

            try:
                # Handle plain text and markdown files
                if ext in [".txt", ".md"]:
                    try:
                        with open(file_path, 'r', encoding='utf-8') as f:
                            content = f.read()
                    except UnicodeDecodeError:
                        with open(file_path, 'r', encoding='latin-1') as f:
                            content = f.read()

                # Handles PDF
                elif ext == ".pdf":
                    content = ""
                    with fitz.open(file_path) as doc:
                        for page in doc:
                            content += page.get_text("text")

                    doc.close()

                # Handles DOCX
                elif ext == ".docx":
                    doc = docx.Document(str(file_path))
                    # Combine all paragraph texts into a single string
                    content = "\n".join([para.text for para in doc.paragraphs])

                else:
                    logger.warning(f"Unknown extension: {ext}.")
                    continue  # Skip unknown file types

                # Add the document to the list with metadata
                documents.append(FileDocument(content, {'source': file_path}))

            except Exception as e:
                # Log an error if reading fails
                logger.error(f"Failed to read {file_path}: {e}")

    logger.info("Document ingestion completed successfully.")

    return documents


def chunk_documents(documents):
    """
    Splits a list of FileDocument objects into smaller, manageable chunks using a RecursiveCharacterTextSplitter.

    This method iterates through each FileDocument, extracts its content, and then uses a
    RecursiveCharacterTextSplitter to break down the content into chunks of a specified size
    with a defined overlap. Empty or whitespace-only chunks are filtered out.
    Each generated chunk is then wrapped back into a FileDocument object, inheriting the
    original document's metadata.

    param:
        documents (list[FileDocument]): A list of FileDocument objects, each containing
                                       'content' (str) and 'metadata' (dict).

    Returns:
        list[FileDocument]: A new list of FileDocument objects, where each object
                            represents a chunk of the original documents' content.
    """

    logger.debug(f"Chunking the documents: {documents}.")

    splitter = RecursiveCharacterTextSplitter(chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP)

    chunks = []

    for doc in documents:
        for chunk in splitter.split_text(doc.content):
            if chunk.strip():  # skip empty or whitespace-only chunks
                chunks.append(FileDocument(chunk, doc.metadata))

    logger.info(f"Chunking completed for the documents: {documents}")

    return chunks
